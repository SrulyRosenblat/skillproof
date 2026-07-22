"""Deterministic checks for the styled status-report .docx.

Independently re-derives the expected section labels, closing line, and
theme colors/fonts from the fixed task inputs (notes.txt, guidelines/*.md,
themes/harbor-blue.yaml) and asserts the produced document matches -- it
never inspects the oracle's or agent's source, only the .docx artifact.
"""
import pathlib
import re

import yaml
from docx import Document
from docx.shared import Pt, RGBColor

WORKSPACE = pathlib.Path("/workspace")
OUTPUT_PATH = WORKSPACE / "output" / "update.docx"


def parse_notes(path):
    categories = []
    current = None
    for raw_line in path.read_text().splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("-"):
            if current is not None:
                current[1].append(line[1:].strip())
        elif line.endswith(":"):
            current = [line[:-1].strip(), []]
            categories.append(current)
    return categories


def parse_guideline(path):
    text = path.read_text()
    lines = text.splitlines()

    intro_lines = []
    for line in lines[1:]:
        if line.startswith("## "):
            break
        intro_lines.append(line)
    intro = " ".join(l.strip() for l in intro_lines if l.strip()).lower()

    sections = []
    if "## Sections" in text:
        block = text.split("## Sections", 1)[1].split("\n## ", 1)[0]
        sections = [m.group(1).strip() for m in re.finditer(r"^\d+\.\s+(.+)$", block, re.MULTILINE)]

    closing = None
    if "## Closing line" in text:
        block = text.split("## Closing line", 1)[1]
        m = re.search(r"^>\s?(.+)$", block, re.MULTILINE)
        if m:
            closing = m.group(1).strip()

    return intro, sections, closing


def select_guideline(notes_categories, guideline_dir):
    header_phrases = [header.lower() for header, _ in notes_categories]
    matches = []
    for path in sorted(guideline_dir.glob("*.md")):
        intro, sections, closing = parse_guideline(path)
        if all(phrase in intro for phrase in header_phrases):
            matches.append((sections, closing))
    assert len(matches) == 1
    return matches[0]


def expected():
    notes_categories = parse_notes(WORKSPACE / "notes.txt")
    section_labels, closing = select_guideline(notes_categories, WORKSPACE / "guidelines")
    theme = yaml.safe_load((WORKSPACE / "themes" / "harbor-blue.yaml").read_text())
    items_by_label = {
        label: items for (_, items), label in zip(notes_categories, section_labels)
    }
    return {
        "section_labels": section_labels,
        "items_by_label": items_by_label,
        "closing": closing,
        "heading_font": theme["heading"]["font_name"],
        "heading_color": RGBColor.from_string(theme["heading"]["font_color_hex"]),
        "heading_size": Pt(theme["heading"]["font_size_pt"]),
        "heading_bold": theme["heading"]["font_bold"],
        "body_font": theme["body"]["font_name"],
        "body_color": RGBColor.from_string(theme["body"]["font_color_hex"]),
        "body_size": Pt(theme["body"]["font_size_pt"]),
        "body_bold": theme["body"]["font_bold"],
    }


EXP = expected()


def _all_runs_styled(paragraph, font_name, color, size, bold):
    runs = paragraph.runs
    assert runs, f"paragraph {paragraph.text!r} has no runs"
    for run in runs:
        assert run.font.name == font_name, (
            f"paragraph {paragraph.text!r}: run font {run.font.name!r} != {font_name!r}"
        )
        assert run.font.color and run.font.color.rgb == color, (
            f"paragraph {paragraph.text!r}: run color {getattr(run.font.color, 'rgb', None)} != {color}"
        )
        assert run.font.size == size, (
            f"paragraph {paragraph.text!r}: run size {run.font.size!r} != {size!r}"
        )
        assert run.font.bold is bold, (
            f"paragraph {paragraph.text!r}: run bold {run.font.bold!r} != {bold!r}"
        )


def test_output_exists():
    assert OUTPUT_PATH.exists(), f"missing {OUTPUT_PATH}"


def test_heading_labels_exact_and_ordered():
    doc = Document(str(OUTPUT_PATH))
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    actual = [p.text.strip() for p in headings]
    assert actual == EXP["section_labels"], (actual, EXP["section_labels"])


def test_headings_use_theme_heading_style():
    doc = Document(str(OUTPUT_PATH))
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    for p in headings:
        _all_runs_styled(p, EXP["heading_font"], EXP["heading_color"], EXP["heading_size"], EXP["heading_bold"])


def test_body_items_present_under_correct_heading_and_styled():
    doc = Document(str(OUTPUT_PATH))
    paras = doc.paragraphs
    heading_idx = {
        i: p.text.strip()
        for i, p in enumerate(paras)
        if p.style.name.startswith("Heading") and p.text.strip() in EXP["section_labels"]
    }
    ordered_positions = sorted(heading_idx)

    for pos_i, start in enumerate(ordered_positions):
        label = heading_idx[start]
        end = ordered_positions[pos_i + 1] if pos_i + 1 < len(ordered_positions) else len(paras)
        body_texts = [paras[j].text.strip() for j in range(start + 1, end) if paras[j].text.strip()]
        expected_items = EXP["items_by_label"][label]
        for item in expected_items:
            assert item in body_texts, f"missing {item!r} under heading {label!r}: {body_texts}"
        for j in range(start + 1, end):
            if paras[j].text.strip():
                _all_runs_styled(paras[j], EXP["body_font"], EXP["body_color"], EXP["body_size"], EXP["body_bold"])


def test_closing_line_verbatim_and_styled():
    doc = Document(str(OUTPUT_PATH))
    non_empty = [p for p in doc.paragraphs if p.text.strip()]
    assert non_empty, "document has no content"
    last = non_empty[-1]
    assert last.text.strip() == EXP["closing"], (last.text.strip(), EXP["closing"])
    _all_runs_styled(last, EXP["body_font"], EXP["body_color"], EXP["body_size"], EXP["body_bold"])
