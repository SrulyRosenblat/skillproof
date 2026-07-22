#!/usr/bin/env python3
"""Computes the styled status-report .docx from the task inputs.

Reads /workspace/notes.txt, picks the one matching guide out of
/workspace/guidelines/*.md by comparing each guide's own description against
the notes' category headers, then builds the document using the section
labels and closing line found in that guide, styled with the fonts/colors
found in /workspace/themes/harbor-blue.yaml.
"""
import re
import pathlib

import yaml
from docx import Document
from docx.shared import Pt, RGBColor

WORKSPACE = pathlib.Path("/workspace")


def parse_notes(path):
    categories = []  # [(header_text, [item, ...]), ...]
    current = None
    for raw_line in path.read_text().splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("-"):
            if current is not None:
                current[1].append(line[1:].strip())
        elif line.endswith(":"):
            current = [line[:-1].strip(), []]
            categories.append(current)
    return categories


def parse_guideline(path):
    text = path.read_text()
    lines = text.splitlines()

    # Intro paragraph: everything between the title and the first "## " block.
    intro_lines = []
    for line in lines[1:]:
        if line.startswith("## "):
            break
        intro_lines.append(line)
    intro = " ".join(l.strip() for l in intro_lines if l.strip()).lower()

    sections = []
    if "## Sections" in text:
        block = text.split("## Sections", 1)[1].split("\n## ", 1)[0]
        sections = [m.group(1).strip() for m in re.finditer(r"^\d+\.\s+(.+)$", block, re.MULTILINE)]

    closing = None
    if "## Closing line" in text:
        block = text.split("## Closing line", 1)[1]
        m = re.search(r"^>\s?(.+)$", block, re.MULTILINE)
        if m:
            closing = m.group(1).strip()

    return intro, sections, closing


def select_guideline(notes_categories, guideline_dir):
    header_phrases = [header.lower() for header, _ in notes_categories]
    matches = []
    for path in sorted(guideline_dir.glob("*.md")):
        intro, sections, closing = parse_guideline(path)
        if all(phrase in intro for phrase in header_phrases):
            matches.append((path, sections, closing))
    assert len(matches) == 1, f"expected exactly one matching guideline, found {[str(m[0]) for m in matches]}"
    return matches[0]


def style_run(run, font_name, hex_color, size_pt, bold):
    run.font.name = font_name
    run.font.color.rgb = RGBColor.from_string(hex_color)
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def build():
    notes_categories = parse_notes(WORKSPACE / "notes.txt")
    _, section_labels, closing = select_guideline(notes_categories, WORKSPACE / "guidelines")
    assert len(section_labels) == len(notes_categories), "section/category count mismatch"

    theme = yaml.safe_load((WORKSPACE / "themes" / "harbor-blue.yaml").read_text())
    heading_font = theme["heading"]["font_name"]
    heading_color = theme["heading"]["font_color_hex"]
    heading_size = theme["heading"]["font_size_pt"]
    heading_bold = theme["heading"]["font_bold"]
    body_font = theme["body"]["font_name"]
    body_color = theme["body"]["font_color_hex"]
    body_size = theme["body"]["font_size_pt"]
    body_bold = theme["body"]["font_bold"]

    doc = Document()
    for (_, items), label in zip(notes_categories, section_labels):
        heading = doc.add_heading(label, level=1)
        for run in heading.runs:
            style_run(run, heading_font, heading_color, heading_size, heading_bold)
        for item in items:
            para = doc.add_paragraph(item, style="List Bullet")
            for run in para.runs:
                style_run(run, body_font, body_color, body_size, body_bold)

    if closing:
        para = doc.add_paragraph(closing)
        for run in para.runs:
            style_run(run, body_font, body_color, body_size, body_bold)

    out_dir = WORKSPACE / "output"
    out_dir.mkdir(parents=True, exist_ok=True)
    doc.save(out_dir / "update.docx")


if __name__ == "__main__":
    build()
